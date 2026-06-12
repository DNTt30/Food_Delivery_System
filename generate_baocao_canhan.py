import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ══════════════════════════════════════════════════════════
# HELPERS dùng chung
# ══════════════════════════════════════════════════════════
def make_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin    = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin   = Cm(3.0)
        s.right_margin  = Cm(2.0)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(13)
    return doc

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_borders(cell, color='AAAAAA'):
    tc       = cell._tc
    tcPr     = tc.get_or_add_tcPr()
    tcBorders= OxmlElement('w:tcBorders')
    for edge in ['top','left','bottom','right']:
        el = OxmlElement('w:%s' % edge)
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_divider(doc, color='1F4E79'):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr= OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1');    bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def banner(doc, text, bg='1F4E79', fg=(0xFF,0xFF,0xFF), size=14):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.rows[0].cells[0]
    c.text = ''
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('  ' + text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = True; r.font.color.rgb = RGBColor(*fg)
    set_cell_bg(c, bg); set_borders(c, bg)

def heading(doc, text, level=1, color=(0x1F,0x4E,0x79)):
    h = doc.add_heading('', level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(16 if level==1 else 14 if level==2 else 13)
    run.bold = True
    return h

def para(doc, text, bold=False, italic=False, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    return p

def bullet(doc, text, bold_part=None, size=13, indent=0.4, color=None):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_before= Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        if text[:idx]:
            rb = p.add_run(text[:idx]); rb.font.name='Times New Roman'; rb.font.size=Pt(size)
        rb2 = p.add_run(bold_part); rb2.font.name='Times New Roman'; rb2.font.size=Pt(size); rb2.bold=True
        if color: rb2.font.color.rgb=RGBColor(*color)
        rest_text = text[idx+len(bold_part):]
        if rest_text:
            ra = p.add_run(rest_text); ra.font.name='Times New Roman'; ra.font.size=Pt(size)
    else:
        r=p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    return p

def issue_table(doc, rows_data, accent='1F4E79'):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style='Table Grid'
    # header
    row0 = tbl.rows[0]
    for i,h in enumerate(['Vấn đề gặp phải','Cách giải quyết']):
        c=row0.cells[i]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.name='Times New Roman'; r.font.size=Pt(11)
        r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        set_cell_bg(c,accent); set_borders(c,accent)
    for idx,(vd,gp) in enumerate(rows_data):
        row=tbl.add_row()
        for i,v in enumerate([vd,gp]):
            c=row.cells[i]; c.text=''
            r=c.paragraphs[0].add_run(v)
            r.font.name='Times New Roman'; r.font.size=Pt(11)
            set_borders(c)
            if idx%2==0: set_cell_bg(c,'F5F5F5')
    return tbl

def cover_page(doc, name, mssv, role, role_desc, accent):
    doc.add_paragraph()
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN')
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.bold=True
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('KHOA CÔNG NGHỆ PHẦN MỀM')
    r.font.name='Times New Roman'; r.font.size=Pt(13); r.bold=True
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)

    doc.add_paragraph()
    doc.add_paragraph()

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('BÁO CÁO CÔNG VIỆC CÁ NHÂN')
    r.font.name='Times New Roman'; r.font.size=Pt(22); r.bold=True
    r.font.color.rgb=RGBColor(*accent)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('Hệ thống Quản lý Giao đồ ăn Trực tuyến')
    r.font.name='Times New Roman'; r.font.size=Pt(16); r.bold=True
    r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('(Food Delivery System)')
    r.font.name='Times New Roman'; r.font.size=Pt(14); r.italic=True

    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    ct=doc.add_table(rows=5,cols=2); ct.style='Table Grid'
    info=[
        ('Họ và tên:', name),
        ('MSSV:',      mssv),
        ('Vai trò:',   role),
        ('Mô tả:',     role_desc),
        ('Ngày nộp:',  'Tháng 6 / 2026'),
    ]
    for ri,(k,v) in enumerate(info):
        row=ct.rows[ri]
        row.cells[0].text=k; row.cells[1].text=v
        for cell in row.cells:
            for rn in cell.paragraphs[0].runs:
                rn.font.name='Times New Roman'; rn.font.size=Pt(12)
        for rn in row.cells[0].paragraphs[0].runs: rn.bold=True
        set_borders(row.cells[0]); set_borders(row.cells[1])

    doc.add_paragraph()
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('TP. Hồ Chí Minh, tháng 6 năm 2026')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.italic=True
    doc.add_page_break()


# ══════════════════════════════════════════════════════════
# BÁO CÁO 1 — DƯƠNG NGỌC TÚ
# ══════════════════════════════════════════════════════════
doc1 = make_doc()
BLUE = (0x1F,0x4E,0x79)

cover_page(doc1,
    name='Dương Ngọc Tú', mssv='22010052',
    role='Team Lead / Fullstack Developer',
    role_desc='Kiến trúc hệ thống, Real-time, Bảo mật, Geocoding, Order Core',
    accent=BLUE)

heading(doc1, 'I. TỔNG QUAN VAI TRÒ', 1, BLUE)
para(doc1,
    'Trong dự án Food Delivery System, em đảm nhận vai trò Nhóm trưởng kiêm Fullstack '
    'Developer. Em chịu trách nhiệm thiết kế kiến trúc tổng thể của hệ thống, xây dựng '
    'các tính năng cốt lõi yêu cầu sự kết hợp chặt chẽ giữa Backend và Frontend, đồng '
    'thời điều phối tiến độ công việc của cả nhóm trong suốt 7 tuần phát triển.')
doc1.add_paragraph()

heading(doc1, 'II. CÁC CÔNG VIỆC ĐÃ THỰC HIỆN', 1, BLUE)

# 2.1
heading(doc1, '2.1. Thiết kế Kiến trúc Hệ thống', 2, BLUE)
para(doc1,
    'Khởi tạo Spring Boot project với phiên bản 3.3.0, Java 17. Thiết lập cấu trúc thư '
    'mục theo mô hình 3 tầng chuẩn: Controller – Service – Repository. Phân chia hệ thống '
    'thành 6 module độc lập: Authentication, Profile, Order, Tracking, Chat, Notification. '
    'Soạn thảo quy chuẩn coding convention và hướng dẫn commit cho toàn nhóm.')

# 2.2
heading(doc1, '2.2. Hệ thống Chat Real-time (WebSocket / STOMP)', 2, BLUE)
para(doc1,
    'Phân tích và quyết định chuyển đổi từ kiến trúc HTTP Polling sang WebSocket STOMP '
    'nhằm giảm độ trễ tin nhắn xuống dưới mili-giây. Xây dựng WebSocketChatController '
    'định tuyến tin nhắn theo vai trò người dùng (Khách hàng, Nhà hàng, Tài xế) và '
    'trạng thái đơn hàng. Triển khai cơ chế tự động khóa luồng chat khi đơn hàng '
    'hoàn thành hoặc bị hủy, đảm bảo tính riêng tư cho tất cả các bên.')
for item in [
    'WebSocketChatController.java — định tuyến STOMP message',
    'ChatService.java — kiểm soát quy tắc chat theo trạng thái đơn',
    'ContactService.java — xác định danh sách liên lạc theo vai trò',
    'PhoneMaskUtil.java — ẩn số điện thoại dạng 098****321 khi đơn đóng',
]:
    bullet(doc1, '• ' + item, size=12)

# 2.3
heading(doc1, '2.3. Core Đặt hàng & Tính phí (Order Core)', 2, BLUE)
para(doc1,
    'Xây dựng OrderService quản lý toàn bộ vòng đời đơn hàng từ lúc khởi tạo '
    '(PENDING) đến khi hoàn thành (COMPLETED). Viết thuật toán Haversine trong '
    'ShippingCalculationService để tính khoảng cách giữa nhà hàng và địa chỉ giao hàng, '
    'từ đó tự động tính phí ship (15.000đ cơ bản + 5.000đ/km thêm) và ước tính ETA. '
    'Xây dựng logic áp dụng Voucher giảm giá dạng phần trăm hoặc số tiền cố định.')
for item in [
    'OrderService.java — logic trung tâm, xử lý transaction toàn bộ đơn hàng',
    'ShippingCalculationService.java — Haversine distance, phí ship, ETA',
    'GeocodingService.java — Nominatim OSM geocoding địa chỉ → tọa độ',
]:
    bullet(doc1, '• ' + item, size=12)

# 2.4
heading(doc1, '2.4. Bảo mật & Định danh Nâng cao', 2, BLUE)
para(doc1,
    'Thực hiện migration toàn bộ khóa chính của bảng User từ kiểu Long sang UUID '
    '(String 36 ký tự) để ngăn chặn tấn công IDOR (Insecure Direct Object Reference). '
    'Tích hợp hệ thống gửi mã OTP xác thực 6 chữ số qua Email sử dụng JavaMailSender '
    'kết nối Gmail SMTP cho quá trình đăng ký và reset mật khẩu. Cấu hình JWT Stateless '
    'không lưu session phía server, tăng khả năng mở rộng hệ thống.')
for item in [
    'JwtUtil.java — tạo và xác thực JWT token',
    'JwtAuthenticationFilter.java — filter kiểm tra token mỗi request',
    'SecurityConfig.java — cấu hình route guards theo role',
    'EmailService.java — gửi OTP qua Gmail SMTP',
]:
    bullet(doc1, '• ' + item, size=12)

# 2.5
heading(doc1, '2.5. Tối ưu Geocoding & Bản đồ', 2, BLUE)
para(doc1,
    'Phân tích nguyên nhân lỗi 403 Forbidden từ Nominatim API và khắc phục bằng cách '
    'bổ sung header User-Agent hợp lệ "FoodDeliveryApp/1.0" vào RestTemplate theo đúng '
    'policy của OpenStreetMap. Phát triển cơ chế Snapshot tọa độ khi đặt đơn và Fallback '
    'tự động cho các đơn hàng cũ thiếu tọa độ. Tích hợp Leaflet Map Picker cho phép '
    'khách hàng và nhà hàng chọn vị trí trực tiếp trên bản đồ.')
doc1.add_paragraph()

heading(doc1, 'III. VẤN ĐỀ GẶP PHẢI & CÁCH GIẢI QUYẾT', 1, BLUE)
issue_table(doc1, [
    ('Tin nhắn Chat có độ trễ cao do dùng HTTP Polling — '
     'client phải gửi request liên tục mỗi 2 giây.',
     'Chuyển sang WebSocket STOMP với In-memory Broker. '
     'Tin nhắn được đẩy tức thời từ server, độ trễ giảm xuống dưới 50ms.'),
    ('Nominatim API trả về lỗi 403 Forbidden — '
     'OpenStreetMap từ chối kết nối do thiếu User-Agent.',
     'Bổ sung header User-Agent = "FoodDeliveryApp/1.0" '
     'vào RestTemplate theo đúng Nominatim Usage Policy.'),
    ('Circular Dependency giữa OrderService và WebSocketConfig '
     'gây lỗi khởi động Spring context.',
     'Tạm thời dùng @Lazy annotation trong WebSocketAuthInterceptor. '
     'Hướng refactor: dùng Spring ApplicationEventPublisher.'),
    ('Bảo mật ID người dùng — Long ID dễ bị đoán, '
     'có nguy cơ tấn công IDOR.',
     'Migration toàn bộ User.id từ Long sang UUID String (36 ký tự), '
     'cập nhật tất cả foreign key liên quan.'),
], accent='1F4E79')
doc1.add_paragraph()

heading(doc1, 'IV. KẾT QUẢ & TỰ ĐÁNH GIÁ', 1, BLUE)
para(doc1,
    'Em đã hoàn thành đầy đủ toàn bộ phần việc được giao, bao gồm kiến trúc hệ thống, '
    'hệ thống chat real-time, xử lý đơn hàng cốt lõi, bảo mật nâng cao và tích hợp bản đồ. '
    'Các tính năng hoạt động ổn định, đã được kiểm thử trên môi trường thực tế với MySQL Aiven Cloud.')
para(doc1,
    'Qua dự án này, em học được cách thiết kế hệ thống phức tạp từ đầu, xử lý các vấn đề '
    'kỹ thuật thực tế (race condition, memory leak, API rate limiting) và điều phối làm việc nhóm '
    'hiệu quả. Em đánh giá mức độ hoàn thành công việc của bản thân đạt 100% theo kế hoạch đề ra.',
    italic=True)

out1 = r'd:\review SPRING_BOOT\springBoot_template-main\Food_Delivery_System\Documents\BAOCAO_DUONGNGOCTU.docx'
os.makedirs(os.path.dirname(out1), exist_ok=True)
doc1.save(out1)
print('✅ Saved: ' + out1)


# ══════════════════════════════════════════════════════════
# BÁO CÁO 2 — ĐINH THỊ NHƯ QUỲNH
# ══════════════════════════════════════════════════════════
doc2 = make_doc()
RED = (0xC0,0x00,0x00)

cover_page(doc2,
    name='Đinh Thị Như Quỳnh', mssv='23010844',
    role='Frontend Developer / Documentation',
    role_desc='UI/UX Design, Thanh toán VNPay/MoMo, Dashboard, Tài liệu dự án',
    accent=RED)

heading(doc2, 'I. TỔNG QUAN VAI TRÒ', 1, RED)
para(doc2,
    'Trong dự án Food Delivery System, em đảm nhận vai trò Frontend Developer và '
    'chuyên viên tài liệu. Em chịu trách nhiệm thiết kế toàn bộ giao diện người dùng '
    'theo phong cách hiện đại, tích hợp cổng thanh toán trực tuyến VNPay và MoMo, '
    'xây dựng các Dashboard quản lý trực quan, và soạn thảo tài liệu dự án.')
doc2.add_paragraph()

heading(doc2, 'II. CÁC CÔNG VIỆC ĐÃ THỰC HIỆN', 1, RED)

heading(doc2, '2.1. Thiết kế UI/UX Toàn hệ thống', 2, RED)
para(doc2,
    'Xây dựng hệ thống giao diện theo phong cách Glassmorphism với Dark Mode, '
    'sử dụng CSS3 Animations và Micro-animations để tạo trải nghiệm người dùng mượt mà, '
    'hiện đại. Đặc biệt tối ưu giao diện theo chuẩn Mobile-First cho màn hình ứng dụng '
    'Tài xế, đảm bảo thao tác dễ dàng trên điện thoại di động khi đang giao hàng.')
for item in [
    'style.css — hệ thống màu sắc HSL, gradient, glassmorphism card, hover effects',
    'customer_layout.html — layout nền thống nhất cho toàn bộ trang khách hàng',
    'navbar_customer.html — thanh điều hướng với notification bell động',
]:
    bullet(doc2, '• ' + item, size=12)

heading(doc2, '2.2. Tích hợp Thanh toán Trực tuyến (VNPay & MoMo)', 2, RED)
para(doc2,
    'Nghiên cứu và tích hợp API của hai cổng thanh toán phổ biến nhất Việt Nam: VNPay '
    'và MoMo vào luồng đặt hàng của hệ thống. Xây dựng giao diện chọn phương thức '
    'thanh toán trực quan, thân thiện người dùng. Xử lý luồng callback từ cổng thanh toán '
    'để cập nhật tự động trạng thái đơn hàng sau khi người dùng hoàn tất giao dịch. '
    'Thiết kế màn hình thông báo kết quả thanh toán thành công và thất bại.')
for item in [
    'Trang chọn phương thức thanh toán: COD / VNPay / MoMo',
    'Xử lý redirect URL và webhook callback từ VNPay, MoMo Sandbox',
    'Màn hình xác nhận thanh toán với thông tin giao dịch chi tiết',
    'Logic rollback đơn hàng khi thanh toán thất bại hoặc hết thời gian chờ',
]:
    bullet(doc2, '• ' + item, size=12)

heading(doc2, '2.3. Module Khách hàng (Customer Frontend)', 2, RED)
para(doc2,
    'Phát triển màn hình Khám phá (Video Review Feed) cho phép khách hàng lướt xem '
    'video ngắn giới thiệu món ăn với thao tác vuốt và thả tim tương tự TikTok/Reels. '
    'Xây dựng Giỏ hàng Global Real-time tự động cập nhật số lượng mà không cần reload '
    'trang. Thiết kế Thanh trạng thái tiến trình đơn hàng (Tracking Progress UI Bar) '
    'hiển thị trực quan các cột mốc: Chuẩn bị – Đang giao – Hoàn thành.')
for item in [
    'Video Review Feed — lướt vuốt, thả tim, lưu món yêu thích',
    'Giỏ hàng toàn cục real-time — cập nhật badge số lượng tức thời',
    'Tracking Progress UI Bar — thanh tiến trình động kèm animation',
    'Danh sách nhà hàng với bộ lọc và tìm kiếm thông minh',
]:
    bullet(doc2, '• ' + item, size=12)

heading(doc2, '2.4. Dashboard & Biểu đồ Thống kê', 2, RED)
para(doc2,
    'Xây dựng bảng Kanban Board cho tài khoản Nhà hàng, cho phép kéo-thả đơn hàng '
    'qua các cột trạng thái: Mới / Đang chuẩn bị / Sẵn sàng giao. Sử dụng thư viện '
    'Chart.js để vẽ biểu đồ thống kê doanh thu theo ngày, tuần, tháng cho Nhà hàng '
    'và Quản trị viên. Thiết kế giao diện thẻ dọc tối ưu cho tài xế trên điện thoại.')
for item in [
    'Kanban Board — kéo thả trạng thái đơn hàng theo cột',
    'Chart.js — biểu đồ đường & cột thống kê doanh thu theo thời gian',
    'Leaflet Map Picker — tích hợp bản đồ chọn vị trí nhà hàng và khách',
    'Admin Dashboard — quản lý tổng quan người dùng, đối tác, voucher',
]:
    bullet(doc2, '• ' + item, size=12)

heading(doc2, '2.5. Tài liệu Dự án (Documentation)', 2, RED)
para(doc2,
    'Soạn thảo và chuẩn hóa toàn bộ tài liệu kỹ thuật và hướng dẫn sử dụng của dự án. '
    'Vẽ các biểu đồ UML bằng Draw.io bao gồm Class Diagram và Use Case Diagram cập nhật '
    'theo đúng code thực tế. Duy trì và cập nhật file README.md, PROJECT_STATE.md và '
    'các tài liệu liên quan trong suốt quá trình phát triển.')
doc2.add_paragraph()

heading(doc2, 'III. VẤN ĐỀ GẶP PHẢI & CÁCH GIẢI QUYẾT', 1, RED)
issue_table(doc2, [
    ('Thymeleaf render trắng trang bản đồ — engine nhầm mảng '
     'JavaScript [[...]] thành inline expression.',
     'Thêm khoảng trắng vào cú pháp: "[ [...] ]" để Thymeleaf '
     'không xử lý như expression, bản đồ hiển thị bình thường.'),
    ('Callback VNPay/MoMo không ổn định — đôi khi webhook '
     'không về kịp trước khi user refresh trang.',
     'Thiết kế màn hình trung gian kiểm tra trạng thái đơn hàng '
     'định kỳ 3 giây, đồng thời hiển thị spinner chờ xác nhận.'),
    ('Giỏ hàng Global không đồng bộ khi mở nhiều tab trình duyệt.',
     'Dùng localStorage làm điểm đồng bộ dữ liệu giỏ hàng, '
     'kết hợp sự kiện window storage event để cập nhật tức thì '
     'khi có thay đổi từ tab khác.'),
    ('Thanh Progress Bar tracking không biết trạng thái hiện tại '
     'của đơn hàng khi vừa load trang.',
     'Gọi API lấy trạng thái đơn ngay khi trang load, '
     'animate thanh progress đến đúng bước hiện tại.'),
], accent='C00000')
doc2.add_paragraph()

heading(doc2, 'IV. KẾT QUẢ & TỰ ĐÁNH GIÁ', 1, RED)
para(doc2,
    'Em đã hoàn thành toàn bộ phần giao diện của hệ thống bao gồm 4 vai trò người dùng, '
    'tích hợp thành công cổng thanh toán VNPay và MoMo ở môi trường Sandbox, '
    'xây dựng hệ thống tài liệu đầy đủ và cập nhật liên tục. '
    'Giao diện được đánh giá hiện đại, thân thiện và hoạt động tốt trên cả desktop lẫn mobile.')
para(doc2,
    'Qua dự án này, em học được kỹ năng thiết kế UI/UX chuyên nghiệp, cách tích hợp '
    'API thanh toán bên thứ ba và viết tài liệu kỹ thuật chuẩn. '
    'Em tự đánh giá mức độ hoàn thành công việc của bản thân đạt 100% theo kế hoạch.',
    italic=True)

out2 = r'd:\review SPRING_BOOT\springBoot_template-main\Food_Delivery_System\Documents\BAOCAO_DINHTHINHUQUYNH.docx'
doc2.save(out2)
print('✅ Saved: ' + out2)


# ══════════════════════════════════════════════════════════
# BÁO CÁO 3 — NGÔ MINH QUÂN
# ══════════════════════════════════════════════════════════
doc3 = make_doc()
GREEN = (0x1E,0x6B,0x2E)

cover_page(doc3,
    name='Ngô Minh Quân', mssv='23017112',
    role='Backend Developer / Database',
    role_desc='Cơ sở dữ liệu, API Backend, Review Reply, Notification, Tracking',
    accent=GREEN)

heading(doc3, 'I. TỔNG QUAN VAI TRÒ', 1, GREEN)
para(doc3,
    'Trong dự án Food Delivery System, em đảm nhận vai trò Backend Developer và '
    'quản trị Cơ sở dữ liệu. Em chịu trách nhiệm thiết kế schema MySQL, xây dựng '
    'các API RESTful cho hệ thống hồ sơ người dùng, quản lý thực đơn, hệ thống '
    'thông báo, phân hệ đánh giá và phản hồi, đồng thời tối ưu hiệu năng truy vấn '
    'và đảm bảo tính ổn định của lớp dữ liệu.')
doc3.add_paragraph()

heading(doc3, 'II. CÁC CÔNG VIỆC ĐÃ THỰC HIỆN', 1, GREEN)

heading(doc3, '2.1. Thiết kế Cơ sở dữ liệu (MySQL / JPA)', 2, GREEN)
para(doc3,
    'Thiết kế và xây dựng toàn bộ mô hình dữ liệu quan hệ cho hệ thống gồm 14 entity '
    'chính và 9 kiểu enum. Cấu hình các mối quan hệ JPA: @OneToOne, @OneToMany, '
    '@ManyToOne với FetchType.LAZY để tránh N+1 Query. Đảm bảo đồng bộ schema '
    'giữa code Entity và MySQL trên môi trường Aiven Cloud.')
for item in [
    'User, CustomerProfile, RestaurantProfile, DriverProfile — hệ thống hồ sơ đa vai trò',
    'FoodOrder, OrderItem, MenuItem, Category — lõi đặt hàng',
    'Payment, Review, Voucher — thanh toán, đánh giá, khuyến mãi',
    'ChatMessage, Notification, OrderTrackingLocation — giao tiếp & tracking',
]:
    bullet(doc3, '• ' + item, size=12)

heading(doc3, '2.2. Module Phản hồi Đánh giá (Review & Reply)', 2, GREEN)
para(doc3,
    'Xây dựng toàn bộ hệ thống đánh giá 2 chiều: Khách hàng gửi đánh giá sao (1–5) '
    'kèm bình luận sau khi đơn hàng hoàn thành, Nhà hàng xem và trực tiếp phản hồi '
    '(reply) lại từng đánh giá. Triển khai logic tự động tính và cập nhật averageRating '
    'của nhà hàng sau mỗi lần có đánh giá mới. Đảm bảo ràng buộc: mỗi đơn hàng '
    'chỉ được đánh giá một lần và chỉ khi trạng thái là COMPLETED.')
for item in [
    'ReviewApiController.java — API gửi đánh giá, lấy danh sách, reply',
    'ReviewService.java — logic kiểm soát, tính averageRating',
    'Review.java (Entity) — lưu rating, comment, reply, createdAt',
    'Ràng buộc: 1 đơn hàng = 1 đánh giá, chỉ khi COMPLETED',
]:
    bullet(doc3, '• ' + item, size=12)

heading(doc3, '2.3. Module Hồ sơ & Phân quyền (Profile / Auth API)', 2, GREEN)
para(doc3,
    'Xây dựng ProfileApiController cung cấp đầy đủ API CRUD cho 4 vai trò: Admin, '
    'Customer, Restaurant và Driver. Mỗi vai trò có endpoint riêng, được bảo vệ bằng '
    'Spring Security theo role. Phối hợp với Tú để cấu hình JwtAuthenticationFilter '
    'đảm bảo phân quyền đúng theo từng route.')
for item in [
    'GET/PUT /api/profile — lấy và cập nhật hồ sơ cá nhân',
    'GET /api/profile/all/:role — lấy danh sách theo vai trò (Admin)',
    'GET/PUT /api/restaurant/profile — hồ sơ riêng cho nhà hàng',
    'Phân quyền Route: CUSTOMER / RESTAURANT / DRIVER / ADMIN',
]:
    bullet(doc3, '• ' + item, size=12)

heading(doc3, '2.4. Module Quản lý Thực đơn & Top Bán Chạy', 2, GREEN)
para(doc3,
    'Xây dựng đầy đủ API CRUD cho thực đơn của nhà hàng: thêm, sửa, xóa, '
    'lấy danh sách món ăn kèm phân trang phía server (Server-side Pagination). '
    'Triển khai logic tự động tăng trường soldCount mỗi khi một đơn hàng chứa '
    'món đó được hoàn thành, phục vụ tính năng "Top Món Bán Chạy" hiển thị trên '
    'trang chủ khách hàng.')
for item in [
    'POST /api/restaurant/menu — thêm món ăn mới kèm upload hình ảnh',
    'PUT /api/restaurant/menu/:id — cập nhật thông tin, giá, trạng thái',
    'DELETE /api/restaurant/menu/:id — xóa món (soft delete)',
    'Logic soldCount — tự động cộng dồn khi đơn hàng COMPLETED',
]:
    bullet(doc3, '• ' + item, size=12)

heading(doc3, '2.5. Hệ thống Thông báo (Notification System)', 2, GREEN)
para(doc3,
    'Phát triển NotificationService với logic Deduplication chống trùng lặp: '
    'tự động gộp các thông báo cùng loại trong khoảng 30 giây thành một bản ghi, '
    'tránh làm spam chuông thông báo khi có nhiều cập nhật liên tục. '
    'Cung cấp đầy đủ API cho frontend: lấy danh sách, đánh dấu đã đọc từng thông báo '
    'và đánh dấu tất cả đã đọc.')
for item in [
    'NotificationService.java — tạo, deduplication, lưu thông báo',
    'NotificationApiController.java — GET list, PUT read, PUT read-all',
    'Notification.java (Entity) — 13 loại NotificationType khác nhau',
    'Hỗ trợ: đơn hàng, thanh toán, chat mới, đánh giá, khuyến mãi...',
]:
    bullet(doc3, '• ' + item, size=12)

heading(doc3, '2.6. Module Tracking Vị trí Tài xế', 2, GREEN)
para(doc3,
    'Xây dựng API nhận tọa độ GPS từ Tài xế theo định kỳ và phát tán tức thời '
    'đến màn hình theo dõi của Khách hàng qua STOMP. Lưu trữ toàn bộ lịch sử '
    'di chuyển bất đồng bộ (@Async) vào bảng order_tracking_locations để tránh '
    'block thread chính, đảm bảo hiệu năng khi có nhiều tài xế hoạt động cùng lúc.')
for item in [
    'POST /api/driver/location — nhận GPS, validate quyền sở hữu đơn',
    'LocationTrackingService.java — @Async lưu lịch sử di chuyển',
    'OrderTrackingLocation.java — entity lưu tọa độ theo TrackingPhase',
    'Validate: Tài xế chỉ được cập nhật GPS của đơn đang giao của mình',
]:
    bullet(doc3, '• ' + item, size=12)
doc3.add_paragraph()

heading(doc3, 'III. VẤN ĐỀ GẶP PHẢI & CÁCH GIẢI QUYẾT', 1, GREEN)
issue_table(doc3, [
    ('N+1 Query gây chậm khi lấy danh sách đơn hàng kèm '
     'chi tiết món — mỗi đơn phát sinh thêm N query con.',
     'Chuyển FetchType sang LAZY toàn bộ, dùng DTO chỉ lấy '
     'đúng trường cần thiết, kết hợp @Query JOIN FETCH '
     'cho các màn hình cần dữ liệu đầy đủ.'),
    ('Schema JPA không đồng bộ với MySQL Aiven Cloud khi '
     'khởi động — lỗi Hibernate validation exception.',
     'Cấu hình ddl-auto=update, rà soát annotation @Column '
     '(length, unique, nullable), đồng bộ lại toàn bộ '
     'constraint giữa code và database.'),
    ('Thông báo bị spam — mỗi lần cập nhật trạng thái '
     'đơn hàng tạo ra nhiều thông báo giống nhau liên tiếp.',
     'Phát triển Deduplication Engine: kiểm tra nếu đã có '
     'thông báo cùng loại trong 30 giây gần nhất thì '
     'cập nhật thay vì tạo mới.'),
    ('Lưu GPS đồng bộ làm chậm response API location update '
     'khi nhiều tài xế gửi tọa độ cùng lúc.',
     'Chuyển LocationTrackingService sang @Async — '
     'API trả về 200 OK ngay lập tức, việc lưu database '
     'thực hiện trên thread pool riêng biệt.'),
], accent='1E6B2E')
doc3.add_paragraph()

heading(doc3, 'IV. KẾT QUẢ & TỰ ĐÁNH GIÁ', 1, GREEN)
para(doc3,
    'Em đã hoàn thành toàn bộ phần Backend và Database được giao, bao gồm thiết kế '
    '14 entity JPA, xây dựng đầy đủ API cho hệ thống hồ sơ, thực đơn, thông báo, '
    'đánh giá và tracking. Hệ thống hoạt động ổn định trên MySQL Aiven Cloud, '
    'không gặp lỗi runtime trong quá trình demo.')
para(doc3,
    'Qua dự án này, em học được cách thiết kế database quan hệ phức tạp, tối ưu hóa '
    'truy vấn JPA và xây dựng các tính năng backend nâng cao như deduplication, async '
    'processing. Em tự đánh giá mức độ hoàn thành công việc của bản thân đạt 100% '
    'theo kế hoạch được giao.',
    italic=True)

out3 = r'd:\review SPRING_BOOT\springBoot_template-main\Food_Delivery_System\Documents\BAOCAO_NGOMINHQUAN.docx'
doc3.save(out3)
print('✅ Saved: ' + out3)

print()
print('🎉 Đã tạo xong 3 file báo cáo cá nhân!')
